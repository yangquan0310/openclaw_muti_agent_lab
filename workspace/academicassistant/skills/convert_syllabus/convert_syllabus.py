#!/usr/bin/env python3
import os
import sys
import subprocess
from docx import Document
import re

def read_docx(file_path):
    """读取docx文件内容"""
    doc = Document(file_path)
    full_text = []
    
    # 读取段落
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # 读取表格
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                full_text.append(" | ".join(row_text))
    
    return "\n\n".join(full_text)

def read_doc(file_path):
    """读取doc文件内容（使用antiword）"""
    try:
        result = subprocess.run(['antiword', file_path], 
                              capture_output=True, text=True, encoding='utf-8')
        if result.returncode == 0:
            return result.stdout
        else:
            print(f"    警告: antiword 返回错误码 {result.returncode}")
            return ""
    except Exception as e:
        print(f"    错误: 读取doc文件失败: {e}")
        return ""

def read_pdf(file_path):
    """读取pdf文件内容（使用pdftotext）"""
    try:
        result = subprocess.run(['pdftotext', file_path, '-'], 
                              capture_output=True, text=True, encoding='utf-8')
        if result.returncode == 0:
            return result.stdout
        else:
            print(f"    警告: pdftotext 返回错误码 {result.returncode}")
            return ""
    except Exception as e:
        print(f"    错误: 读取pdf文件失败: {e}")
        return ""

def convert_to_markdown(text):
    """将文本转换为Markdown格式"""
    lines = text.split('\n')
    markdown_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 标题检测
        if re.match(r'^一、', line) or re.match(r'^二、', line) or re.match(r'^三、', line) or \
           re.match(r'^四、', line) or re.match(r'^五、', line) or re.match(r'^六、', line) or \
           re.match(r'^七、', line) or re.match(r'^八、', line) or re.match(r'^九、', line) or \
           re.match(r'^十、', line):
            markdown_lines.append(f"## {line}")
        elif re.match(r'^（一）', line) or re.match(r'^（二）', line) or re.match(r'^（三）', line) or \
             re.match(r'^（四）', line) or re.match(r'^（五）', line):
            markdown_lines.append(f"### {line}")
        elif re.match(r'^1\.', line) or re.match(r'^2\.', line) or re.match(r'^3\.', line):
            markdown_lines.append(f"#### {line}")
        else:
            # 表格行
            if '|' in line:
                markdown_lines.append(line)
            else:
                markdown_lines.append(line)
    
    return '\n\n'.join(markdown_lines)

def process_project(project_path):
    """处理单个项目"""
    project_name = os.path.basename(project_path)
    print(f"处理项目: {project_name}")
    
    original_doc_dir = os.path.join(project_path, '原始文档')
    syllabus_dir = os.path.join(project_path, '课程大纲')
    
    if not os.path.exists(original_doc_dir):
        print(f"  跳过: 原始文档目录不存在")
        return False
    
    if not os.path.exists(syllabus_dir):
        os.makedirs(syllabus_dir)
        print(f"  创建课程大纲目录")
    
    # 查找原始文档中的文件
    all_files = os.listdir(original_doc_dir)
    supported_files = [f for f in all_files if f.endswith(('.docx', '.doc', '.pdf'))]
    
    if not supported_files:
        print(f"  跳过: 没有找到支持的文件（支持docx/doc/pdf）")
        return False
    
    # 处理每个文件
    for src_file in supported_files:
        src_path = os.path.join(original_doc_dir, src_file)
        print(f"  读取: {src_file}")
        
        try:
            # 根据文件类型读取内容
            if src_file.endswith('.docx'):
                content = read_docx(src_path)
            elif src_file.endswith('.doc'):
                content = read_doc(src_path)
            elif src_file.endswith('.pdf'):
                content = read_pdf(src_path)
            else:
                continue
            
            if not content or not content.strip():
                print(f"    警告: 文件内容为空")
                continue
            
            # 转换为Markdown
            markdown_content = convert_to_markdown(content)
            
            # 生成输出文件名
            base_name = os.path.splitext(src_file)[0]
            md_file = f"{base_name}.md"
            md_path = os.path.join(syllabus_dir, md_file)
            
            # 写入Markdown文件
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(f"# {project_name} 教学大纲\n\n")
                f.write(markdown_content)
            
            print(f"  写入: {md_file}")
            
        except Exception as e:
            print(f"  错误: 处理 {src_file} 时出错: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    return True

def main():
    base_path = os.path.expanduser('~/教研室仓库/教务归档/项目文件')
    
    if not os.path.exists(base_path):
        print(f"错误: 项目文件目录不存在: {base_path}")
        sys.exit(1)
    
    # 获取所有项目目录
    projects = [d for d in os.listdir(base_path) 
                if os.path.isdir(os.path.join(base_path, d)) 
                and not d.startswith('.')]
    
    print(f"找到 {len(projects)} 个项目\n")
    
    success_count = 0
    for project in projects:
        project_path = os.path.join(base_path, project)
        if process_project(project_path):
            success_count += 1
        print()
    
    print(f"完成: 成功处理 {success_count}/{len(projects)} 个项目")

if __name__ == '__main__':
    main()
