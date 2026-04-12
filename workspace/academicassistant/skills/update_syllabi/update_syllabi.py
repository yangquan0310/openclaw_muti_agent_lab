#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
教学大纲更新脚本
重新读取各个项目原始文件中的教学大纲，并且更新到课程大纲中
"""

import os
from docx import Document
import re
import pdfplumber
import subprocess

# 项目文件根目录
PROJECTS_ROOT = os.path.expanduser("~/教研室仓库/教务归档/项目文件/")

def doc_to_markdown(doc_path):
    """
    将旧版Word文档(.doc)转换为Markdown格式
    """
    try:
        # 使用antiword提取文本
        result = subprocess.run(
            ['antiword', doc_path],
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='replace'
        )
        if result.returncode == 0:
            return result.stdout
        else:
            print(f"    antiword警告: {result.stderr}")
            return ""
    except Exception as e:
        print(f"    处理.doc文件时出错: {str(e)}")
        return ""

def pdf_to_markdown(pdf_path):
    """
    将PDF文档转换为Markdown格式
    """
    markdown_lines = []
    
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                markdown_lines.append(text)
                markdown_lines.append("")
            # 尝试提取表格
            tables = page.extract_tables()
            for table in tables:
                if table and len(table) > 0:
                    markdown_lines.append("")
                    # 表头
                    header_cells = table[0]
                    header = "| " + " | ".join(str(cell).strip() if cell else "" for cell in header_cells) + " |"
                    markdown_lines.append(header)
                    # 分隔线
                    separator = "| " + " | ".join("---" for _ in header_cells) + " |"
                    markdown_lines.append(separator)
                    # 数据行
                    for row in table[1:]:
                        row_text = "| " + " | ".join(str(cell).strip() if cell else "" for cell in row) + " |"
                        markdown_lines.append(row_text)
                    markdown_lines.append("")
    
    return "\n".join(markdown_lines)

def docx_to_markdown(docx_path):
    """
    将Word文档转换为Markdown格式
    """
    doc = Document(docx_path)
    markdown_lines = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            markdown_lines.append("")
            continue
        
        # 处理标题
        if para.style.name.startswith('Heading') or para.style.name.startswith('标题'):
            level = 1
            # 尝试从样式名称中提取级别
            match = re.search(r'(\d+)', para.style.name)
            if match:
                level = int(match.group(1))
            # 限制标题级别
            level = min(level, 6)
            markdown_lines.append(f"{'#' * level} {text}")
        else:
            markdown_lines.append(text)
    
    # 处理表格
    for table in doc.tables:
        markdown_lines.append("")
        # 表头
        if len(table.rows) > 0:
            header_cells = table.rows[0].cells
            header = "| " + " | ".join(cell.text.strip() for cell in header_cells) + " |"
            markdown_lines.append(header)
            # 分隔线
            separator = "| " + " | ".join("---" for _ in header_cells) + " |"
            markdown_lines.append(separator)
            # 数据行
            for row in table.rows[1:]:
                row_cells = row.cells
                row_text = "| " + " | ".join(cell.text.strip() for cell in row_cells) + " |"
                markdown_lines.append(row_text)
        markdown_lines.append("")
    
    return "\n".join(markdown_lines)

def update_project_syllabus(project_dir):
    """
    更新单个项目的教学大纲
    """
    project_name = os.path.basename(project_dir)
    print(f"\n处理项目: {project_name}")
    
    # 构建路径
    original_doc_dir = os.path.join(project_dir, "原始文档")
    syllabus_dir = os.path.join(project_dir, "课程大纲")
    
    if not os.path.exists(original_doc_dir):
        print(f"  警告: 原始文档目录不存在: {original_doc_dir}")
        return False
    
    if not os.path.exists(syllabus_dir):
        print(f"  创建课程大纲目录: {syllabus_dir}")
        os.makedirs(syllabus_dir, exist_ok=True)
    
    # 查找原始文档中的文件
    all_files = os.listdir(original_doc_dir)
    docx_files = [f for f in all_files if f.endswith('.docx')]
    pdf_files = [f for f in all_files if f.endswith('.pdf')]
    doc_files = [f for f in all_files if f.endswith('.doc')]
    
    if not docx_files and not pdf_files and not doc_files:
        print(f"  警告: 未找到.docx/.pdf/.doc文件在: {original_doc_dir}")
        return False
    
    success_count = 0
    
    # 处理.docx文件
    for docx_file in docx_files:
        docx_path = os.path.join(original_doc_dir, docx_file)
        print(f"  读取: {docx_file}")
        
        try:
            # 转换为markdown
            markdown_content = docx_to_markdown(docx_path)
            
            # 生成输出文件名
            base_name = os.path.splitext(docx_file)[0]
            md_file = f"{base_name}.md"
            md_path = os.path.join(syllabus_dir, md_file)
            
            # 写入markdown文件
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            print(f"  ✓ 已更新: {md_file}")
            success_count += 1
            
        except Exception as e:
            print(f"  ✗ 处理失败 {docx_file}: {str(e)}")
    
    # 处理.pdf文件
    for pdf_file in pdf_files:
        pdf_path = os.path.join(original_doc_dir, pdf_file)
        print(f"  读取: {pdf_file}")
        
        try:
            # 转换为markdown
            markdown_content = pdf_to_markdown(pdf_path)
            
            # 生成输出文件名
            base_name = os.path.splitext(pdf_file)[0]
            md_file = f"{base_name}.md"
            md_path = os.path.join(syllabus_dir, md_file)
            
            # 写入markdown文件
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            print(f"  ✓ 已更新: {md_file}")
            success_count += 1
            
        except Exception as e:
            print(f"  ✗ 处理失败 {pdf_file}: {str(e)}")
    
    # 处理.doc文件
    for doc_file in doc_files:
        doc_path = os.path.join(original_doc_dir, doc_file)
        print(f"  读取: {doc_file}")
        
        try:
            # 使用antiword转换为markdown
            markdown_content = doc_to_markdown(doc_path)
            
            if markdown_content:
                # 生成输出文件名
                base_name = os.path.splitext(doc_file)[0]
                md_file = f"{base_name}.md"
                md_path = os.path.join(syllabus_dir, md_file)
                
                # 写入markdown文件
                with open(md_path, 'w', encoding='utf-8') as f:
                    f.write(markdown_content)
                
                print(f"  ✓ 已更新: {md_file}")
                success_count += 1
            else:
                print(f"  ⚠  未能从 {doc_file} 提取内容")
            
        except Exception as e:
            print(f"  ✗ 处理失败 {doc_file}: {str(e)}")
    
    return success_count > 0

def main():
    """
    主函数
    """
    print("=" * 60)
    print("教学大纲更新脚本")
    print("=" * 60)
    
    if not os.path.exists(PROJECTS_ROOT):
        print(f"错误: 项目根目录不存在: {PROJECTS_ROOT}")
        return
    
    # 获取所有项目目录
    project_dirs = [
        os.path.join(PROJECTS_ROOT, d) 
        for d in os.listdir(PROJECTS_ROOT) 
        if os.path.isdir(os.path.join(PROJECTS_ROOT, d))
    ]
    
    print(f"找到 {len(project_dirs)} 个项目")
    
    success_count = 0
    for project_dir in sorted(project_dirs):
        if update_project_syllabus(project_dir):
            success_count += 1
    
    print("\n" + "=" * 60)
    print(f"完成! 成功处理 {success_count}/{len(project_dirs)} 个项目")
    print("=" * 60)

if __name__ == "__main__":
    main()
