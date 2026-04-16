#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX文档生成器 v2.0 (优化版)
用于创建包含页脚的Microsoft Word文档

优化内容：
1. 增强参数验证和错误处理
2. 添加更多样式控制功能
3. 支持图片插入
4. 支持代码块、引用等样式
5. 改进表格功能
6. 更好的类型提示
7. 支持链式调用
"""

import os
import re
from typing import Optional, List, Dict, Union, Tuple, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class DocxGeneratorError(Exception):
    """DOCX生成器异常基类"""
    pass


class ValidationError(DocxGeneratorError):
    """参数验证错误"""
    pass


# 样式预设
STYLE_PRESETS = {
    'title': {
        'size': 18,
        'bold': True,
        'color': '000000',
        'alignment': 'center'
    },
    'subtitle': {
        'size': 14,
        'bold': True,
        'color': '333333',
        'alignment': 'left'
    },
    'normal': {
        'size': 11,
        'bold': False,
        'color': '000000',
        'alignment': 'left'
    },
    'code': {
        'size': 9,
        'font': 'Consolas',
        'color': '006400',
        'bg_color': 'F5F5F5'
    },
    'quote': {
        'size': 11,
        'italic': True,
        'color': '666666',
        'indent': 0.5
    },
    'warning': {
        'size': 11,
        'bold': True,
        'color': '8B0000'
    },
    'info': {
        'size': 11,
        'bold': True,
        'color': '00008B'
    }
}


class DocxGenerator:
    """
    DOCX文档生成器类（优化版）

    功能：
    - 创建包含页脚的Word文档
    - 丰富的样式控制
    - 图片、代码块、引用等支持
    - 链式调用支持
    - 完善的错误处理
    """

    def __init__(self, page_size: str = 'A4'):
        """
        初始化文档生成器

        Args:
            page_size: 页面大小 ('A4', 'Letter')
        """
        self.doc = Document()
        self.footer_text = "AI生成，仅供参考"
        self.header_text: Optional[str] = None
        self._setup_page(page_size)

    def _setup_page(self, page_size: str):
        """设置页面大小"""
        if page_size == 'A4':
            self.doc.sections[0].page_height = Cm(29.7)
            self.doc.sections[0].page_width = Cm(21)
        elif page_size == 'Letter':
            self.doc.sections[0].page_height = Inches(11)
            self.doc.sections[0].page_width = Inches(8.5)

    # ==================== 文档设置 ====================

    def set_footer_text(self, text: str):
        """设置页脚文本"""
        if not isinstance(text, str):
            raise ValidationError("页脚文本必须是字符串")
        self.footer_text = text
        return self  # 支持链式调用

    def set_header_text(self, text: str):
        """设置页眉文本"""
        if not isinstance(text, str):
            raise ValidationError("页眉文本必须是字符串")
        self.header_text = text
        return self

    # ==================== 内容添加 ====================

    def add_title(self, title: str, level: int = 1):
        """
        添加标题

        Args:
            title: 标题文本
            level: 标题级别 (1-9)

        Returns:
            self (支持链式调用)
        """
        if not title:
            raise ValidationError("标题不能为空")
        if not 1 <= level <= 9:
            raise ValidationError("标题级别必须在1-9之间")

        self.doc.add_heading(title, level=level)
        return self

    def add_paragraph(
        self,
        text: str,
        style: Optional[str] = None,
        font_size: Optional[int] = None,
        bold: bool = False,
        italic: bool = False,
        color: Optional[str] = None,
        alignment: Optional[str] = None
    ):
        """
        添加段落（支持样式设置）

        Args:
            text: 段落文本
            style: 段落样式名称
            font_size: 字体大小（磅）
            bold: 是否加粗
            italic: 是否斜体
            color: 字体颜色（十六进制，如 'FF0000'）
            alignment: 对齐方式 ('left', 'center', 'right', 'justify')

        Returns:
            段落对象
        """
        if text is None:
            text = ""

        para = self.doc.add_paragraph(text)

        if style:
            try:
                para.style = style
            except KeyError:
                pass

        # 应用格式
        if font_size or bold or italic or color or alignment:
            run = para.runs[0] if para.runs else para.add_run()
            if font_size:
                run.font.size = Pt(font_size)
            if bold:
                run.font.bold = True
            if italic:
                run.font.italic = True
            if color:
                run.font.color.rgb = self._hex_to_rgb(color)

            if alignment:
                para.alignment = self._get_alignment(alignment)

        return self  # 支持链式调用

    def add_list(
        self,
        items: List[str],
        ordered: bool = False,
        level: int = 0
    ):
        """
        添加列表

        Args:
            items: 列表项内容
            ordered: 是否为有序列表
            level: 列表层级 (0-8)

        Returns:
            self
        """
        if not items:
            raise ValidationError("列表项不能为空")

        if not 0 <= level <= 8:
            raise ValidationError("列表层级必须在0-8之间")

        # python-docx 使用固定的列表样式名称
        style = 'List Number' if ordered else 'List Bullet'

        for item in items:
            self.doc.add_paragraph(str(item), style=style)

        return self

    def add_table(
        self,
        data: List[List[str]],
        style: str = 'Table Grid',
        header_row: bool = True,
        autofit: bool = True
    ):
        """
        添加表格（增强版）

        Args:
            data: 二维列表，表示表格数据
            style: 表格样式
            header_row: 第一行是否为表头（加粗）
            autofit: 是否自动调整列宽

        Returns:
            表格对象
        """
        if not data:
            raise ValidationError("表格数据不能为空")

        if not data[0]:
            raise ValidationError("表格第一行不能为空")

        # 创建表格
        rows = len(data)
        cols = len(data[0])

        # 验证所有行列数一致
        for row in data:
            if len(row) != cols:
                raise ValidationError("表格数据行列数不一致")

        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = style

        # 填充数据
        for i, row_data in enumerate(data):
            for j, cell_data in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = str(cell_data)

                # 表头加粗
                if header_row and i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

        # 自动调整列宽
        if autofit:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    if paragraphs:
                        paragraphs[0].runs[0].font.size = Pt(10)

        return table

    def add_code_block(self, code: str, language: str = 'Python'):
        """
        添加代码块

        Args:
            code: 代码内容
            language: 编程语言名称（用于注释）

        Returns:
            self
        """
        if not code:
            raise ValidationError("代码内容不能为空")

        # 添加代码段落
        para = self.doc.add_paragraph()
        run = para.add_run(code)

        # 设置代码样式
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 100, 0)  # 深绿色

        # 设置段落格式
        para_format = para.paragraph_format
        para_format.left_indent = Inches(0.5)
        para_format.right_indent = Inches(0.5)
        para_format.space_before = Pt(6)
        para_format.space_after = Pt(6)

        # 添加浅灰色背景（使用边框模拟）
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'F5F5F5')
        para._element.get_or_add_pPr().append(shading_elm)

        return self

    def add_quote(self, text: str):
        """
        添加引用块

        Args:
            text: 引用文本

        Returns:
            self
        """
        if not text:
            raise ValidationError("引用文本不能为空")

        para = self.doc.add_paragraph()
        run = para.add_run(f'"{text}"')

        # 设置引用样式
        run.font.italic = True
        run.font.color.rgb = RGBColor(102, 102, 102)
        run.font.size = Pt(11)

        # 添加缩进
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.right_indent = Inches(0.5)

        return self

    def add_image(
        self,
        image_path: str,
        width: Optional[float] = None,
        height: Optional[float] = None,
        align: str = 'center'
    ):
        """
        添加图片

        Args:
            image_path: 图片路径
            width: 宽度（英寸）
            height: 高度（英寸）
            align: 对齐方式 ('left', 'center', 'right')

        Returns:
            self
        """
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"图片文件不存在: {image_path}")

        # 添加图片
        para = self.doc.add_paragraph()
        run = para.add_run()

        # 设置图片大小
        width_kwargs = {}
        if width:
            width_kwargs['width'] = Inches(width)
        if height:
            width_kwargs['height'] = Inches(height)

        run.add_picture(image_path, **width_kwargs)

        # 设置对齐
        para.alignment = self._get_alignment(align)

        return self

    def add_hyperlink(self, text: str, url: str):
        """
        添加超链接

        Args:
            text: 链接文本
            url: 链接地址

        Returns:
            self
        """
        if not text or not url:
            raise ValidationError("链接文本和URL不能为空")

        # 创建超链接
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.font.underline = True

        # 添加超链接关系
        run.hyperlink = url

        return self

    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()
        return self

    def add_horizontal_rule(self):
        """添加水平线"""
        para = self.doc.add_paragraph()
        run = para.add_run('_' * 80)
        run.font.color.rgb = RGBColor(200, 200, 200)
        return self

    def add_spacing(self, lines: int = 1):
        """
        添加空行

        Args:
            lines: 空行数

        Returns:
            self
        """
        for _ in range(lines):
            self.doc.add_paragraph()
        return self

    # ==================== 辅助方法 ====================

    def _get_alignment(self, alignment: str) -> WD_PARAGRAPH_ALIGNMENT:
        """获取对齐方式枚举"""
        alignment_map = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
            'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        }
        return alignment_map.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)

    def _hex_to_rgb(self, hex_color: str) -> RGBColor:
        """将十六进制颜色转换为RGB"""
        hex_color = hex_color.lstrip('#')
        if len(hex_color) != 6:
            raise ValidationError("颜色代码必须是6位十六进制")

        # 验证是否为有效的十六进制
        if not all(c in '0123456789ABCDEFabcdef' for c in hex_color):
            raise ValidationError(f"无效的颜色代码: {hex_color}")

        try:
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
        except ValueError:
            raise ValidationError(f"无效的颜色代码: {hex_color}")

        return RGBColor(r, g, b)

    # ==================== 页眉页脚 ====================

    def _add_header(self):
        """添加页眉"""
        if not self.header_text:
            return

        section = self.doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]

        header_para.text = self.header_text
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 设置页眉字体样式
        run = header_para.runs[0]
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.italic = True

    def _add_footer(self):
        """添加页脚"""
        section = self.doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]

        footer_para.text = self.footer_text
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 设置页脚字体样式
        if footer_para.runs:
            run = footer_para.runs[0]
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.italic = True

    # ==================== 保存文档 ====================

    def save(self, filepath: str) -> str:
        """
        保存文档

        Args:
            filepath: 保存路径

        Returns:
            保存的文件路径
        """
        try:
            # 验证文件路径
            if not filepath:
                raise ValidationError("文件路径不能为空")

            # 自动添加 .docx 扩展名
            if not filepath.endswith('.docx'):
                filepath += '.docx'

            # 添加页眉页脚
            self._add_header()
            self._add_footer()

            # 确保目录存在
            dir_path = os.path.dirname(filepath)
            if dir_path:
                os.makedirs(dir_path, exist_ok=True)

            # 保存文档
            self.doc.save(filepath)

            return filepath

        except PermissionError:
            raise DocxGeneratorError(f"没有权限写入文件: {filepath}")
        except Exception as e:
            raise DocxGeneratorError(f"保存文档失败: {str(e)}")


# ==================== 便捷函数 ====================

def create_simple_document(
    title: str,
    content: str,
    output_path: str,
    footer_text: str = "AI生成，仅供参考",
    header_text: Optional[str] = None
) -> str:
    """
    创建简单文档（便捷函数）

    Args:
        title: 文档标题
        content: 文档内容
        output_path: 输出路径
        footer_text: 页脚文本
        header_text: 页眉文本

    Returns:
        保存的文件路径
    """
    try:
        generator = DocxGenerator()
        if header_text:
            generator.set_header_text(header_text)
        generator.set_footer_text(footer_text)
        generator.add_title(title)
        generator.add_paragraph(content)
        return generator.save(output_path)
    except Exception as e:
        raise DocxGeneratorError(f"创建文档失败: {str(e)}")


def create_report(
    title: str,
    sections: List[Dict[str, Any]],
    output_path: str,
    footer_text: str = "AI生成，仅供参考",
    header_text: Optional[str] = None
) -> str:
    """
    创建报告文档（便捷函数，增强版）

    Args:
        title: 报告标题
        sections: 章节列表，每个章节可包含：
            - title: 章节标题
            - content: 章节内容
            - type: 内容类型 ('paragraph', 'list', 'code', 'quote')
            - items: 列表项（当 type='list' 时）
            - code: 代码内容（当 type='code' 时）
        output_path: 输出路径
        footer_text: 页脚文本
        header_text: 页眉文本

    Returns:
        保存的文件路径
    """
    try:
        generator = DocxGenerator()
        if header_text:
            generator.set_header_text(header_text)
        generator.set_footer_text(footer_text)

        generator.add_title(title, level=1)

        for section in sections:
            section_type = section.get('type', 'paragraph')

            if section_type == 'paragraph':
                if 'title' in section:
                    generator.add_title(section['title'], level=2)
                if 'content' in section:
                    generator.add_paragraph(section['content'])

            elif section_type == 'list':
                if 'title' in section:
                    generator.add_title(section['title'], level=2)
                if 'items' in section:
                    generator.add_list(
                        section['items'],
                        ordered=section.get('ordered', False)
                    )

            elif section_type == 'code':
                if 'title' in section:
                    generator.add_title(section['title'], level=2)
                if 'code' in section:
                    generator.add_code_block(section['code'])

            elif section_type == 'quote':
                if 'text' in section:
                    generator.add_quote(section['text'])

            elif section_type == 'table':
                if 'title' in section:
                    generator.add_title(section['title'], level=2)
                if 'data' in section:
                    generator.add_table(section['data'])

        return generator.save(output_path)

    except Exception as e:
        raise DocxGeneratorError(f"创建报告失败: {str(e)}")


# ==================== 示例代码 ====================

def main():
    """示例：创建文档"""
    print("正在创建示例文档（v2.0）...")

    try:
        # 示例1: 简单文档
        create_simple_document(
            title="测试文档 v2.0",
            content="这是优化版DOCX生成器创建的文档。",
            output_path="./output/v2_simple_test.docx"
        )

        # 示例2: 复杂文档（使用链式调用）
        gen = DocxGenerator()

        (gen
            .add_title("项目报告 v2.0", level=1)
            .add_paragraph("本报告使用优化版生成器创建。")
            .add_spacing(1))

        # 添加章节
        gen.add_title("功能特性", level=2)
        gen.add_list([
            "增强的参数验证",
            "更多样式控制",
            "图片插入支持",
            "代码块和引用"
        ], ordered=True)

        gen.add_spacing(1)

        # 添加代码块
        gen.add_title("代码示例", level=2)
        gen.add_code_block("def hello():\n    print('Hello, World!'")

        # 添加表格
        gen.add_title("数据对比", level=2)
        gen.add_table([
            ["版本", "功能数", "代码行数"],
            ["v1.0", "10", "224"],
            ["v2.0", "25+", "600+"]
        ])

        gen.add_spacing(1)

        # 添加引用
        gen.add_quote("代码是写给人看的，只是顺便给机器运行。")

        gen.save("./output/v2_complex_test.docx")

        print("✓ 文档创建成功！")
        print("  - v2_simple_test.docx")
        print("  - v2_complex_test.docx")

    except Exception as e:
        print(f"✗ 创建失败: {str(e)}")


if __name__ == "__main__":
    main()
